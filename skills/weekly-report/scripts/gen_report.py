"""
gen_report.py - 生成周报/日报 .docx

用法:
  python gen_report.py --type weekly --date 2026-02-28 --data '<json>'
  python gen_report.py --type daily --date 2026-02-27 --data '<json>'

数据格式 (JSON):
{
  "author": "李晶",
  "bugs": [
    {"id": "46912", "title": "问题描述", "resolvedDate": "02-25"},
    ...
  ],
  "tasks": [
    {"id": "49058", "title": "任务名称", "finishedDate": "02-25"},
    ...
  ],
  "summary": "工作总结文本",
  "next_plan": ["计划1", "计划2"]
}

输出: workspace/reports/YYYY-MM-DD/weekly_report_YYYY-MM-DD.docx
"""
import argparse
import json
import os
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def create_report(report_type, date, data):
    doc = Document()

    # 默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    author = data.get('author', '')
    type_label = '周报' if report_type == 'weekly' else '日报'

    # 标题
    title = doc.add_heading('', level=0)
    run = title.add_run(f'{type_label} — {author}')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 日期
    date_text = data.get('date_range', date)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(date_text)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # 本周/今日工作内容
    period = '本周' if report_type == 'weekly' else '今日'
    doc.add_heading(f'{period}工作内容', level=1)

    if data.get('work_title'):
        doc.add_heading(data['work_title'], level=2)

    if data.get('work_desc'):
        doc.add_paragraph(data['work_desc'])

    # Bug 表格
    bugs = data.get('bugs', [])
    if bugs:
        doc.add_heading(f'解决 Bug（{len(bugs)} 个）', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0].cells
        hdr[0].text = 'Bug ID'
        hdr[1].text = '问题描述'
        hdr[2].text = '解决日期'
        for cell in hdr:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        for bug in bugs:
            row = table.add_row().cells
            row[0].text = f'#{bug["id"]}'
            row[1].text = bug.get('title', '')
            row[2].text = bug.get('resolvedDate', '')
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for cell in row:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9.5)

        doc.add_paragraph()

    # 任务表格
    tasks = data.get('tasks', [])
    if tasks:
        doc.add_heading(f'完成任务（{len(tasks)} 个）', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0].cells
        hdr[0].text = '任务 ID'
        hdr[1].text = '任务名称'
        hdr[2].text = '完成日期'
        for cell in hdr:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        for task in tasks:
            row = table.add_row().cells
            row[0].text = f'#{task["id"]}'
            row[1].text = task.get('title', '')
            row[2].text = task.get('finishedDate', '')
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for cell in row:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9.5)

        doc.add_paragraph()

    # 工作总结
    summary = data.get('summary', '')
    if summary:
        doc.add_heading('工作总结', level=1)
        doc.add_paragraph(summary)

    # 下周/明日计划
    next_plan = data.get('next_plan', [])
    if next_plan:
        next_label = '下周计划' if report_type == 'weekly' else '明日计划'
        doc.add_heading(next_label, level=1)
        for item in next_plan:
            doc.add_paragraph(item, style='List Bullet')

    # 输出目录: workspace/reports/YYYY-MM-DD/
    script_dir = os.path.dirname(os.path.abspath(__file__))
    workspace_dir = os.path.normpath(os.path.join(script_dir, '..', '..', '..'))
    report_dir = os.path.join(workspace_dir, 'reports', date)
    os.makedirs(report_dir, exist_ok=True)

    filename = f'{type_label}_{date}.docx'
    output_path = os.path.join(report_dir, filename)
    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    sys.stdout = open(sys.stdout.fileno(), mode='w', encoding='utf-8', closefd=False)

    parser = argparse.ArgumentParser(description='生成周报/日报')
    parser.add_argument('--type', choices=['weekly', 'daily'], default='weekly')
    parser.add_argument('--date', required=True, help='报告日期 YYYY-MM-DD')
    parser.add_argument('--data', help='JSON 数据字符串')
    parser.add_argument('--data-file', help='JSON 数据文件路径')
    args = parser.parse_args()

    if args.data_file:
        with open(args.data_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
    elif args.data:
        data = json.loads(args.data)
    else:
        print('ERROR: 需要 --data 或 --data-file 参数')
        sys.exit(1)
    path = create_report(args.type, args.date, data)
    print(f'OK: {path}')
